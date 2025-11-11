from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import numpy as np
import joblib
from collections import Counter
from underthesea import word_tokenize
import time

# Load dữ liệu gốc
train_data = pd.read_csv('data/train_data.csv')
test_data = pd.read_csv('data/test_data.csv')

# Tạo lại cột Processed_Question
def preprocess_text(text):
    tokens = word_tokenize(text.lower())
    stopwords = ['tôi', 'đang', 'là', 'có', 'không', 'và', 'hoặc', 'nhưng', 'nếu', 'thì']
    tokens = [word for word in tokens if word not in stopwords and word.isalpha()]
    return ' '.join(tokens)

train_data['Processed_Question'] = train_data['Question'].apply(preprocess_text)

# Load dữ liệu đã xử lý
X_train_ml = np.load('data/X_train.npy')
X_test_ml = np.load('data/X_test.npy')
y_train = np.load('data/y_train.npy')
y_test = np.load('data/y_test.npy')

# Load deep learning data
X_train_dl = np.load('data/X_train_dl.npy')
X_test_dl = np.load('data/X_test_dl.npy')

# Load tokenizer và metadata
tokenizer = joblib.load('data/tokenizer.joblib')
label_encoder = joblib.load('data/label_encoder.joblib')
dl_metadata = joblib.load('data/dl_metadata.joblib')

# Tạo document
doc = Document()

# Tiêu đề chính
title = doc.add_heading('ViMedical - Hệ Chuyên Gia Bệnh Tật', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Bước 2: Tiền Xử Lý Dữ Liệu', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Thêm thông tin tổng quan
doc.add_paragraph(f'Ngày tạo báo cáo: {time.strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph('')

# Mục tiêu
doc.add_heading('🎯 Mục Tiêu', 2)
doc.add_paragraph('Chuẩn bị và tiền xử lý dữ liệu văn bản tiếng Việt cho các mô hình học máy và học sâu.')

# Tổng quan dữ liệu
doc.add_heading('📈 Tổng Quan Dữ Liệu', 2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Thống Kê'
hdr_cells[1].text = 'Giá Trị'

rows_data = [
    ('Tổng số mẫu training', str(len(train_data))),
    ('Tổng số mẫu test', str(len(test_data))),
    ('Số lớp bệnh tật', str(len(np.unique(y_train)))),
    ('Số từ trung bình/câu', f"{len(train_data['Question'].str.split().sum())/len(train_data):.1f}")
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# Phân tích văn bản gốc
doc.add_heading('🔤 Phân Tích Văn Bản Gốc', 2)

questions = train_data['Question']
word_counts = questions.str.split().str.len()
char_counts = questions.str.len()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Thống Kê'
hdr_cells[1].text = 'Giá Trị'

text_stats = [
    ('Tổng số câu hỏi', str(len(questions))),
    ('Số từ trung bình/câu', f"{word_counts.mean():.1f}"),
    ('Số từ tối đa/câu', str(word_counts.max())),
    ('Số từ tối thiểu/câu', str(word_counts.min())),
    ('Số ký tự trung bình/câu', f"{char_counts.mean():.1f}")
]

for stat in text_stats:
    row_cells = table.add_row().cells
    row_cells[0].text = stat[0]
    row_cells[1].text = stat[1]

# Tiền xử lý văn bản
doc.add_heading('🧹 Tiền Xử Lý Văn Bản', 2)
doc.add_paragraph('Các bước xử lý:')
doc.add_paragraph('• Chuyển về chữ thường', style='List Bullet')
doc.add_paragraph('• Tokenization bằng Underthesea', style='List Bullet')
doc.add_paragraph('• Loại bỏ stopwords cơ bản', style='List Bullet')
doc.add_paragraph('• Giữ lại chỉ từ alphabetic', style='List Bullet')

# Ví dụ văn bản đã xử lý
doc.add_heading('📝 Ví Dụ Văn Bản Đã Xử Lý', 2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Văn Bản Gốc'
hdr_cells[1].text = 'Văn Bản Đã Xử Lý'

for i in range(min(5, len(train_data))):
    row_cells = table.add_row().cells
    original = train_data['Question'].iloc[i][:80] + "..." if len(train_data['Question'].iloc[i]) > 80 else train_data['Question'].iloc[i]
    processed = str(train_data['Processed_Question'].iloc[i])[:80] + "..." if len(str(train_data['Processed_Question'].iloc[i])) > 80 else str(train_data['Processed_Question'].iloc[i])
    row_cells[0].text = original
    row_cells[1].text = processed

# Chuẩn bị dữ liệu cho các mô hình
doc.add_heading('🤖 Chuẩn Bị Dữ Liệu Cho Các Mô Hình', 2)

# Traditional ML
doc.add_heading('📊 Dữ Liệu Cho Mô Hình Học Máy Truyền Thống', 3)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Thông Số'
hdr_cells[1].text = 'Giá Trị'

ml_stats = [
    ('Kích thước từ vựng', str(X_train_ml.shape[1])),
    ('Số mẫu training', str(X_train_ml.shape[0])),
    ('Số tính năng', str(X_train_ml.shape[1])),
    ('Tỷ lệ sparsity', f"{np.sum(X_train_ml) / X_train_ml.size:.3f}")
]

for stat in ml_stats:
    row_cells = table.add_row().cells
    row_cells[0].text = stat[0]
    row_cells[1].text = stat[1]

# Deep Learning
doc.add_heading('🧠 Dữ Liệu Cho Mô Hình Học Sâu', 3)

vocab_size = dl_metadata['vocab_size']
max_seq_len = dl_metadata['max_sequence_length']

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Thông Số'
hdr_cells[1].text = 'Giá Trị'

dl_stats = [
    ('Kích thước từ vựng', str(vocab_size)),
    ('Số mẫu training', str(X_train_dl.shape[0])),
    ('Độ dài sequence tối đa', str(max_seq_len)),
    ('Tổng số từ duy nhất', str(len(tokenizer.word_index)))
]

for stat in dl_stats:
    row_cells = table.add_row().cells
    row_cells[0].text = stat[0]
    row_cells[1].text = stat[1]

# Ví dụ tokenization
doc.add_heading('🔢 Ví Dụ Tokenization', 3)

sample_text = train_data['Processed_Question'].iloc[0]
sample_sequence = X_train_dl[0][:20]

p = doc.add_paragraph()
p.add_run('Văn bản gốc: ').bold = True
p.add_run(sample_text)

p = doc.add_paragraph()
p.add_run('Sequence (20 token đầu): ').bold = True
p.add_run(str(sample_sequence.tolist()))

p = doc.add_paragraph()
p.add_run('Mapping một số token:').bold = True

word_index = tokenizer.word_index
for word, idx in list(word_index.items())[:10]:
    doc.add_paragraph(f'{word} → {idx}', style='List Bullet')

# Phân tích nhãn
doc.add_heading('🏷️ Phân Tích Nhãn (Labels)', 2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Thông Số'
hdr_cells[1].text = 'Giá Trị'

label_stats = [
    ('Tổng số lớp bệnh', str(len(label_encoder.classes_))),
    ('Lớp có nhiều mẫu nhất', str(max(Counter(y_train).values()))),
    ('Lớp có ít mẫu nhất', str(min(Counter(y_train).values()))),
    ('ID lớp phổ biến nhất', str(Counter(y_train).most_common(1)[0][0]))
]

for stat in label_stats:
    row_cells = table.add_row().cells
    row_cells[0].text = stat[0]
    row_cells[1].text = stat[1]

# Top 10 lớp bệnh
doc.add_heading('📋 Top 10 Lớp Bệnh Phổ Biến Nhất', 3)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Tên Bệnh'
hdr_cells[1].text = 'Số Mẫu'
hdr_cells[2].text = 'Tỷ Lệ (%)'

label_counts = Counter(y_train)
total_samples = len(y_train)
for label_id, count in label_counts.most_common(10):
    disease_name = label_encoder.inverse_transform([label_id])[0]
    percentage = (count / total_samples) * 100
    row_cells = table.add_row().cells
    row_cells[0].text = disease_name
    row_cells[1].text = str(count)
    row_cells[2].text = f"{percentage:.2f}%"

# Tệp đã lưu
doc.add_heading('💾 Tệp Đã Lưu', 2)
files_list = [
    'data/X_train.npy - Features training (Traditional ML)',
    'data/y_train.npy - Labels training',
    'data/X_test.npy - Features test (Traditional ML)',
    'data/y_test.npy - Labels test',
    'data/X_train_dl.npy - Sequences training (Deep Learning)',
    'data/X_test_dl.npy - Sequences test (Deep Learning)',
    'data/tokenizer.joblib - Tokenizer cho Deep Learning',
    'data/label_encoder.joblib - Label Encoder',
    'data/dl_metadata.joblib - Metadata cho Deep Learning'
]

for file_desc in files_list:
    doc.add_paragraph(file_desc, style='List Bullet')

# Tổng kết
doc.add_heading('📝 Tổng Kết', 2)
summary_text = f"""Đã hoàn thành tiền xử lý dữ liệu:
• Xử lý {len(train_data)} mẫu training và {len(test_data)} mẫu test
• Tạo dữ liệu cho {len(np.unique(y_train))} loại bệnh khác nhau
• Chuẩn bị dữ liệu cho cả mô hình học máy truyền thống và học sâu
• Từ vựng: {vocab_size} từ cho Deep Learning, {X_train_ml.shape[1]} tính năng cho Traditional ML
• Sẵn sàng cho bước tiếp theo: Lựa chọn và huấn luyện mô hình"""

doc.add_paragraph(summary_text)

# Lưu file
doc.save('reports/Step2_Data_Preprocessing.docx')
print("✅ Đã tạo báo cáo DOCX: reports/Step2_Data_Preprocessing.docx")